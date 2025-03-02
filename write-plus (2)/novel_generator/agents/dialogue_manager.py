import json
from novel_generator.utils.file_handler import FileHandler
from novel_generator.utils.prompt_generator import PromptGenerator
from docx import Document
from docx.shared import Inches
import os
import logging
import re
from typing import Dict, List, Union
import time
import threading
import queue

import signal
import sys

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class DialogueManager:
    def __init__(self, ai1, ai2, export_frequency=1, output_dir="output"):
        self.ai1 = ai1
        self.ai2 = ai2
        self.export_frequency = export_frequency  # 每多少轮导出设定
        self.turn = 1  # 对话轮次
        self.file_handler = FileHandler()
        self.prompt_generator = PromptGenerator()
        self.exported_settings = {}  # 存储每次导出的设定
        self.output_dir = './output'  # 定义输出目

        self.chapter_content = ""  # 用于保存章节内容，合并后导出word
        self.novel_information ="" # 用于保存小说信息
        self.full_outline = "" # 保存完整的大纲信息
        self.chapter_titles = "" # 保存章节标题
        self.current_state = "continue" # 保存当前状态
        self.bg_settings = {} # 保存小说背景设定
        self.part_chapter = []#这是一个列表，用来存储每一章章节的内容
        self.last_chapter = "" # 保存上一章节的内容
        self.last_commit = "" # 保存上一章节的评价
    def start_dialogue(self, initial_prompt, a1_gen, a2_gen):
        logging.info("开始后续对话...")
        last_chapter = a1_gen
        last_commit = a2_gen
        while self.current_state == "continue":
            self.turn += 1
            # 动态生成 AI1 的提示词，包含大纲、目录和最新章节
            ai1_prompt = self.generate_dynamic_prompt_ai1(last_chapter, last_commit)
            print(len(ai1_prompt))
            # print("--------------------ai1提示词:",ai1_prompt[:20])
            response_ai1 = self.ai1.send_message(ai1_prompt)
            if response_ai1 is None:
                logging.warning("AI1 返回 None, 尝试重新发送")
                response_ai1 = self.ai1.send_message(ai1_prompt) # 尝试重新发送
                if response_ai1 is None:
                   logging.error("AI1 多次返回 None, 对话结束")
                   self.current_state = "stop" # 设置停止状态
                   break

            logging.info(f"AI1: {response_ai1}")
            timeout = 15
            print("ai1文章内容生成完毕,请指示, 如果你对这一章节有意见,请输入任意内容并按回车,本章节将在{}秒后送往编辑处审核:".format(timeout))
            prompt = ""
            interrupet = self.input_with_timeout(prompt, timeout)
            # interrupet = self.get_input_with_timeout("如果你对这一章节有意见,请输入任意内容并按回车,本章节将在{}秒后送往编辑处审核:  ".format(timeout), timeout=timeout)
            if interrupet == "":
                print("本章节已送往编辑处审核")
            else:
                usr_advice = input("请输入修改意见:") 
                advice = f"请根据以下修改意见修改当前章节的小说,并直接给出修改后的内容: {usr_advice}"
                response_ai1 = self.ai1.send_message(advice)
                logging.info(f"AI1: {response_ai1}")
            
            self.save_chapter_content(response_ai1)#保存当前章节

            # 动态生成 AI2 的提示词，包含大纲、目录和最新章节，以及AI1的最新输出
            ai2_prompt = self.generate_dynamic_prompt_ai2(response_ai1, last_chapter)
            last_chapter = response_ai1 #审核完之后把当前章节当作下一章节
            print(len(ai2_prompt))
            print("-------------------------------------------------")
            response_ai2 = self.ai2.send_message(ai2_prompt)
            last_commit = response_ai2
            print("-------------------------------------------------")
            if response_ai2 is None:
                logging.warning("AI2 返回 None, 尝试重新发送")
                response_ai2 = self.ai2.send_message(ai2_prompt)  # 尝试重新发送
                if response_ai2 is None:
                  logging.error("AI2 多次返回 None, 对话结束")
                  self.current_state = "stop" # 设置停止状态
                  break
            
            print("------------------------------------------------")

            logging.info(f"AI2: {response_ai2}")

            # settings = self.extract_settings()  # 提取设置
            # settings["turn"] = self.turn
            # self.export_settings(settings)
            # if self.turn % self.export_frequency == 0:
            #      self.update_prompts_with_settings(settings) # 只在需要的时候更新prompt
            # # 保存当前状态
            # print("---------------存状态分界线----------------------------------")
            # self.save_current_state()
        logging.info("对话结束")
    
    def input_with_timeout(self, prompt, timeout):
        """改进的超时输入函数"""
        user_input = queue.Queue()
        
        def input_thread():
            try:
                text = input(prompt)
                user_input.put(text)
            except:
                user_input.put("")
        
        # 创建线程并在启动前设置daemon状态
        thread = threading.Thread(target=input_thread)
        thread.daemon = True  # 在启动前设置daemon状态
        thread.start()
        
        try:
            # 等待用户输入，最多等待timeout秒
            return user_input.get(timeout=timeout)
        except queue.Empty:
            print(f"\n等待输入超时（{timeout}秒），继续执行...")
            return ""
 
    
    # def get_input_with_timeout(self, prompt, timeout):
    #     """带超时的输入函数"""
    #     user_input = queue.Queue()
        
    #     def input_thread():
    #         try:
    #             text = input(prompt)
    #             user_input.put(text)
    #         except:
    #             user_input.put(None)
    #     thread = threading.Thread(target=input_thread)
    #     thread.daemon = True
    #     thread.start()
    #     try:
    #         # 等待输入，超时时间为timeout秒
    #         return user_input.get(timeout=timeout)
    #     except queue.Empty:
    #         return ""
        
    def generate_dynamic_prompt_ai1(self, a1_gen, a2_gen):
         """动态生成 AI1 的提示词"""
         # 从AI1的初始prompt中提取核心指令和约束
         original_prompt = self.ai1.system_prompt
        #  combined_prompt = f"{original_prompt}\n\n"
         combined_prompt = ""
         # 添加完整的大纲和目录信息
         print("ai1提示词生成中")
         bg_settings_str = f""
         for key, value in self.bg_settings.items():
             bg_settings_str += f"{key}: {value}\n"
         combined_prompt += f"小说创作的基础方向, 包括类型, 主角, 世界背景等:\n{bg_settings_str}\n\n"
         print("小说背景设定已加入")
         combined_prompt += f"小说设定以及大纲目录章节:\n{self.novel_information}\n\n"
         print("章节的标题已加入")
         # 添加已经生成的章节内容
        #  if self.chapter_content:
        #       combined_prompt += f"已生成章节的内容:\n{self.chapter_content}\n\n"
         if a1_gen is not None:
           combined_prompt += f"最新的章节的内容:\n{a1_gen}\n\n"
           print("最新的章节的内容已加入")
         combined_prompt += f"小说编辑的反馈意见:\n{a2_gen}\n\n"
         combined_prompt += f"请根据上一章的设定和内容, 直接给出下一章内容(无需修改原有章节), 确保逻辑连贯, 请使用中文, 并且每一章的字数不少于2000字"
        #  print("--------------a1提示词:",combined_prompt)
         return combined_prompt
    
    def generate_dynamic_prompt_ai2(self, current_chapter, last_chapter):
        """动态生成 AI2 的提示词"""
         # 从AI2的初始prompt中提取核心指令和约束
        original_prompt = self.ai2.system_prompt
        combined_prompt = ""
        # combined_prompt = f"{original_prompt}\n\n"

        # Add combined chat history
        print("ai2提示词生成中")
        bg_settings_str = f""
        for key, value in self.bg_settings.items():
            bg_settings_str += f"{key}: {value}\n"
        combined_prompt += f"小说创作的基础方向, 包括类型, 主角, 世界背景等:\n{bg_settings_str}\n\n"
        print("小说背景设定已加入")
        combined_prompt += f"小说设定以及大纲目录章节:\n{self.novel_information}\n\n"
        print("章节的标题已加入")

        combined_prompt += f"最新的章节的内容:\n{current_chapter}\n\n"
        combined_prompt += f"上一章节的内容:\n{last_chapter}\n\n"
        print("最新的两张章节已加入")
        combined_prompt += f"请审核最新章节内容, 检查逻辑、设定是否符合网文小说以及是否符合小说背景设定, 并审查它跟上一章节内容是否连贯,注意仅需审查当前最新的章节即可"
        return combined_prompt

    def export_settings(self, settings):
        # 保存设定到文件
        self.save_settings(settings)

    def extract_background_settings(self, text):
         """从对话历史中提取小说背景设定并保存"""
         logging.info("正在从对话中提取小说背景设定...")
         parts = text.split('【')
         result = {}
    
         # 处理每个部分
         for part in parts[1:]:  # 跳过第一个空字符串
             name_content = part.split('】')
             if len(name_content) == 2:
                 name = name_content[0]
                 content = name_content[1].strip()
                 result[name] = content
         self.bg_settings = result
         return result

    
    def save_bg_settings(self):
        filename = os.path.join(self.output_dir, "novel_settings.txt")
    
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                for key, value in self.bg_settings.items():
                    f.write(f"【{key}】:\n{value}\n")
                    f.write("-" * 50 + "\n")  # 写入50个连字符作为分隔符
        
            logging.info(f"设定已保存到 {filename}")
        except Exception as e:
            logging.error(f"保存设定失败: {e}")

    def save_novel_index(self):
        filename = os.path.join(self.output_dir, "novel_index.txt")
    
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(f"{self.novel_information}\n")
            logging.info(f"设定已保存到 {filename}")
        except Exception as e:
            logging.error(f"保存设定失败: {e}")

    def save_chapter_content(self,chapter_content ):
        filename = os.path.join(self.output_dir, "chapter_content_{}.txt".format(self.turn))
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(f"{chapter_content}\n")
            logging.info(f"设定已保存到 {filename}")
        except Exception as e:
            logging.error(f"保存设定失败: {e}")

    # def extract_settings(self):
    #      """从对话历史中提取设定"""
    #      logging.info("正在从对话中提取设定...")
    #      ai1_history = self.ai1.get_history()
    #      print("--------------ai1_history:",ai1_history)
    #      ai2_history = self.ai2.get_history()
    #     #  print("--------------ai2_history:",ai2_history)

    #      combined_history = []
    #      extracted_settings = {}

    #      for item in ai1_history:
    #         print("---------开始提取设定：：:")
    #         if item['role'] == 'assistant':
    #             print("---------开始提取设定2")
    #             combined_history.append(f"{item['role']}: {item['content']}")
    #              # 尝试从 AI1 的响应中提取设定
    #             for setting_type in ["**分卷大纲**", "**世界设定**", "**人物设定**", "**前50章标题**"]:
    #                print("--------------尝试提取设定:",setting_type)
    #                if setting_type in item['content']:
    #                    try:
    #                         print("--------------提取设定:",setting_type)
    #                         print("--------------提取设定:",item['content'].split(setting_type)[1].strip())
    #                         extracted_settings[setting_type] = item['content'].split(setting_type)[1].strip()
    #                    except Exception as e:
    #                      logging.error(f"提取 {setting_type} 失败: {e}") 

    #      for item in ai2_history:
    #         if item['role'] == 'assistant':
    #              combined_history.append(f"{item['role']}: {item['content']}")
    #              # 从AI2回复中提取反馈意见
    #              if "修改建议" in item['content']:
    #                  try:
    #                    extracted_settings["修改建议"] = item['content'].split("修改建议")[1].strip()
    #                  except Exception as e:
    #                      logging.error(f"提取 修改建议 失败：{e}")

    #      extracted_settings["history"] = combined_history
    #     #  for key in extracted_settings:
    #     #     print(key, extracted_settings[key])
    #      return extracted_settings

    def save_settings(self, settings):
         """保存设定到文件，文件名包含当前轮数"""
         print("---------保存设定到文件，文件名包含当前轮数！！！！！！！！！！")
         for item in settings.get("history", []):
            if True:
                print("开始提取内容了！！！！！！！！！！")
                if "**大纲**" in item:
                    filename = f"{self.output_dir}/outline_turn_{settings['turn']}.txt"
                    try:
                         content = item.split("**大纲**")[1].strip() # 提取大纲内容，去除多余的空格
                         self.file_handler.save_text_to_file(content, filename)
                         logging.info(f"大纲已保存到 {filename}")
                         self.full_outline = content # 保存完整的大纲信息
                    except Exception as e:
                         logging.error(f"保存大纲失败: {e}")
                elif "**世界设定**" in item:
                     filename = f"{self.output_dir}/world_setting_turn_{settings['turn']}.txt"
                     try:
                         content = item.split("**世界设定**")[1].strip()
                         self.file_handler.save_text_to_file(content, filename)
                         logging.info(f"世界设���已保存到 {filename}")
                     except Exception as e:
                        logging.error(f"保存世界设定失败: {e}")
                elif "**人物设定**" in item:
                    filename = f"{self.output_dir}/character_setting_turn_{settings['turn']}.txt"
                    try:
                        content = item.split("**人物设定**")[1].strip()
                        self.file_handler.save_text_to_file(content, filename)
                        logging.info(f"人物设定已保存到 {filename}")
                    except Exception as e:
                       logging.error(f"保存人物设定失败: {e}")
                elif "**前50章标题**" in item:
                    filename = f"{self.output_dir}/first_fifty_chapter_titles_turn_{settings['turn']}.txt"
                    try:
                        content = item.split("**前50章标题**")[1].strip()
                        self.file_handler.save_text_to_file(content, filename)
                        logging.info(f"前50章标题已保存到 {filename}")
                        self.chapter_titles = content # 保存章节标题
                    except Exception as e:
                        logging.error(f"保存前50章标题失败: {e}")
                elif "**章节**" in item:
                     try:
                        content = item.split("**章节**")[1].strip()
                        if "**本章完**" in content:
                            content = content.split("**本章完**")[0].strip()  # 删除尾部的本章完
                        self.chapter_content += f"\n\n**章节{settings['turn']}**:\n{content}"  # 合并章节
                        self.part_chapter.append(content) #保存单章章节内容
                        print("当前章节字数 content:",len(content))
                        logging.info(f"章节{settings['turn']}内容已保存")
                        self.export_word_document(settings['turn'], content) # 每保存章节就导出
                     except Exception as e:
                        logging.error(f"保存章节内容失败: {e}")
         # Export chat histories
         ai1_chat_filename = f"{self.output_dir}/ai1_chat_history_turn_{settings['turn']}.json"
         self.ai1.export_chat_history(ai1_chat_filename)
         ai2_chat_filename = f"{self.output_dir}/ai2_chat_history_turn_{settings['turn']}.json"
         self.ai2.export_chat_history(ai2_chat_filename)


         self.exported_settings[settings["turn"]] = settings  # 添加到导出列表

    def get_exported_settings(self):
         return self.exported_settings

    def update_prompts_with_settings(self, settings):
        """将提取的设定融入到新的提示词中"""
        logging.info("正在更新提示词")
        ai1_new_prompt = self.prompt_generator.generate_new_prompt(self.ai1.system_prompt, settings)
        ai2_new_prompt = self.prompt_generator.generate_new_prompt(self.ai2.system_prompt, settings)

        self.ai1.messages = [{"role": "system", "content": ai1_new_prompt}]  # 重置ai1的对话历史，放入新的prompt
        self.ai2.messages = [{"role": "system", "content": ai2_new_prompt}]  # 重置ai2的对话历史，放入新的prompt

        logging.info("提示词已更新")
    def export_word_document(self, turn, content):
        """导出word文档"""
        document = Document()
        document.add_heading(f"章节 {turn}", level=1)
        document.add_paragraph(content) # 添加章节内容
        output_filename = os.path.join(self.output_dir, f"chapter_{turn}.docx") # 输出的文件名和路径
        try:
           document.save(output_filename)
           logging.info(f"章节内容已导出到 {output_filename}")
        except Exception as e:
           logging.error(f"导出word文档失败：{e}")

    def save_current_state(self):
        """保存当前状态，包括轮数，大纲和章节内容"""
        state_data = {
           "turn": self.turn,
           "full_outline": self.full_outline,
           "chapter_titles": self.chapter_titles,
           "chapter_content": self.chapter_content,
           "current_state": self.current_state
        }
        filename = os.path.join(self.output_dir, "current_state.json")
        try:
           self.file_handler.save_json_to_file(state_data, filename)
           logging.info(f"当前状态已保存到 {filename}")
        except Exception as e:
           logging.error(f"保存当前状态失败：{e}")
    def load_current_state(self):
        """加载当前状态"""
        filename = os.path.join(self.output_dir, "current_state.json")
        try:
           state_data = self.file_handler.read_json_from_file(filename)
           self.turn = state_data.get("turn", 0)
           self.full_outline = state_data.get("full_outline", "")
           self.chapter_content = state_data.get("chapter_content", "")
           self.current_state = state_data.get("current_state", "continue")
           logging.info(f"已从 {filename} 加载当前状态")
        except Exception as e:
           logging.error(f"加载当前状态失败: {e}")
